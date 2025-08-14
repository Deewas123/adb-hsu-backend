from __future__ import annotations
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt

# --- Simple HSU cleanups ---
UK_US = {
    'colour':'color','labour':'labor','organisation':'organization',
    'programme':'program','defence':'defense','centre':'center'
}
EN_RANGE = re.compile(r'(\d)\s*-\s*(\d)')          # 3 - 5 -> 3–5
EM_TIGHT = re.compile(r'\s*—\s*')                  # spaces around em-dash -> none
OXFORD  = re.compile(r'(\b[^,]+,\s+[^,\s]+)\s+(and|or)\s+([^,\s][^.;:!?)]*)')
DATE_NUM = re.compile(r'\b(\d{1,4})[/-](\d{1,2})[/-](\d{2,4})\b')
MONTHS = ["January","February","March","April","May","June","July","August","September","October","November","December"]

def us_spelling(t):
    def repl(m):
        s=m.group(0); low=s.lower(); rep=UK_US.get(low)
        if not rep: return s
        if s.isupper(): return rep.upper()
        if s.istitle(): return rep.title()
        return rep
    return re.sub(r'\b('+'|'.join(map(re.escape,UK_US.keys()))+r')\b', repl, t, flags=re.I)

def en_dash_ranges(t): return EN_RANGE.sub(r'\1–\2', t)
def tighten_em(t):    return EM_TIGHT.sub('—', t)
def oxford(t):
    def r(m):
        p,c,tail=m.group(1),m.group(2),m.group(3)
        if not p.endswith(','): p+=','
        return f"{p} {c} {tail}"
    return OXFORD.sub(r, t)

def normalize_dates(t):
    def r(m):
        a,b,c=m.groups()
        if len(a)==4: y=int(a); mo=int(b); d=int(c)
        else:
            if int(a)>12: d=int(a); mo=int(b); y=int(c)
            elif int(b)>12: mo=int(a); d=int(b); y=int(c)
            else: d=int(a); mo=int(b); y=int(c)
        if 1<=mo<=12 and 1<=d<=31 and 1900<=y<=2100:
            return f"{d} {MONTHS[mo-1]} {y}"
        return m.group(0)
    return DATE_NUM.sub(r, t)

def apply_rules(p):
    t=p.text
    for fn in (us_spelling, en_dash_ranges, tighten_em, oxford, normalize_dates):
        t=fn(t)
    if t!=p.text: p.text=t

def enforce_fonts(p, size_pt=11):
    for r in p.runs:
        r.font.name="Times New Roman"
        r.font.size=Pt(size_pt)

def process_docx_bytes(data: bytes) -> bytes:
    doc = Document(BytesIO(data))
    # paragraphs
    for p in doc.paragraphs:
        apply_rules(p)
        s=(p.style.name or '').lower() if p.style else ''
        if 'heading 1' in s: enforce_fonts(p,14)
        elif 'heading 2' in s: enforce_fonts(p,12)
        elif 'heading 3' in s: enforce_fonts(p,11)
        elif 'caption' in s:   enforce_fonts(p,10)
        else: enforce_fonts(p,11)
    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for pr in cell.paragraphs:
                    apply_rules(pr); enforce_fonts(pr,9)
    out=BytesIO(); doc.save(out); return out.getvalue()
